
from openai import OpenAI
import yaml
from docx import Document
import re
import pandas as pd

def list_headings(doc_path):
    doc = Document(doc_path)
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            print(f"{para.style.name}: {para.text}")
        else:
            print(para.style.name)

def extract_title(heading):
    pattern = r"[（(]?[一二三四五六七八九十]+[)）]?\s*([^、\n]+)|[一二三四五六七八九十]+、\s*([^、\n]+)|(^\S+篇)"
    match = re.search(pattern, heading)
    if match:
        return match.group(1) or match.group(2) or match.group(3)
    return heading

def merge_content(content):
    merged_content = []
    new_content = ""
    for line in content:
        if line.startswith('问：') and new_content:
            merged_content.append(new_content)
            new_content = ""
        new_content += line
    merged_content.append(new_content)
    return merged_content

def word2csv(doc_path):
    """
    将word转换成csv
    四列分别是：part / section / subsection / qapair
    """
    doc = Document(doc_path)
    structured_doc = pd.DataFrame(columns=['部分', '考试类型', '问题类型', '问答对'])
    part_list = ['高考学考', '中考中招', '研考成考', '自学考试', '证书考试']
    current_h2 = None
    current_h3 = None
    content = []

    part_index = 0
    part = None
    for para in doc.paragraphs:
        if para.style.name == "Heading 2":  # 发现新的 H2
            current_h2 = extract_title(para.text)  # 更新当前大标题
            if para.text.startswith('一、'):
                part = part_list[part_index]
                part_index += 1

        elif para.style.name == "Heading 3":  # 发现新的 H3
            # 储存上一个部分
            if content:
                for qa_pair in merge_content(content):
                    new_row = pd.DataFrame([{'部分': part, '考试类型': current_h2, '问题类型': current_h3, '问答对': qa_pair}])
                    structured_doc = pd.concat([structured_doc, new_row], ignore_index=True)
                content = []
            current_h3 = extract_title(para.text)  # 更新当前小标题

        else:  # 普通文本
            if current_h2 is None or current_h3 is None or para.text.strip()=='':
                continue  # 避免在没有 H2 时存储文本
            content.append(para.text)


    # 处理最后一块内容
    if content:
        for qa_pair in merge_content(content):
            new_row = pd.DataFrame([{'部分': part, '考试类型': current_h2, '问题类型': current_h3, '问答对': qa_pair}])
            structured_doc = pd.concat([structured_doc, new_row], ignore_index=True)
              
    return structured_doc        

if __name__ == '__main__':   
    # 读取 Word 文档的标题
    filepath = '/Users/wanghelin/Documents/毕设/code-langgraph/data/考试院faq第1版第7.00稿（打印版）-修改20250508 - 高考 一到六.docx'
    structured_doc = word2csv(filepath)
    structured_doc.to_csv('/Users/wanghelin/Documents/毕设/code-langgraph/data/build/考试院faq第1版第7.00稿（打印版）-修改20250508 - 高考 一到六.csv', index=False)